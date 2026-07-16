# -*- coding: utf-8 -*-
"""Biblioteca para gerar petições sobre o Timbrado.docx do escritório,
preservando cabeçalho/rodapé (timbrado) e aplicando a formatação padrão:
Calibri Light 12 (corpo) / 10 (citações, recuo 4cm), justificado, espaçamento 1,2,
6pt depois, recuo de primeira linha 1,25cm, títulos em negrito.

O caminho do timbrado é OBRIGATÓRIO e vem do escritorio.json
(escritorios[].estrutura.timbrado) — não há default embutido, porque cada
escritório tem o seu, em pastas diferentes.
"""
import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

FONT = "Calibri Light"


def _setfont(run, size=12, bold=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(a), FONT)


class Peticao:
    def __init__(self, timbrado):
        if not timbrado:
            raise ValueError(
                "Peticao(timbrado=...) exige o caminho do Timbrado.docx. "
                "Pegue-o em escritorios[].estrutura.timbrado do escritorio.json."
            )
        if not os.path.exists(timbrado):
            raise FileNotFoundError(
                f"Timbrado não encontrado: {timbrado}\n"
                "Confira o caminho em escritorio.json ou rode 'peticia configurar'."
            )
        self.doc = Document(timbrado)
        st = self.doc.styles["Normal"]
        st.font.name = FONT
        st.font.size = Pt(12)
        rpr = st.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(a), FONT)

    def _base(self, p):
        pf = p.paragraph_format
        pf.line_spacing = 1.2
        pf.space_after = Pt(6)
        pf.space_before = Pt(0)
        return pf

    def address(self, text):
        p = self.doc.add_paragraph()
        self._base(p)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(0)
        _setfont(p.add_run(text), 12, True)
        return p

    def title(self, text):
        # Centraliza (herança do estilo antigo). A formatação-padrão marketplace
        # exige títulos numerados à ESQUERDA: use subtitle() para os títulos de
        # seção (1., 2., 3.1…), não title().
        p = self.doc.add_paragraph()
        pf = self._base(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_before = Pt(6)
        _setfont(p.add_run(text), 12, True)
        return p

    def subtitle(self, text):
        p = self.doc.add_paragraph()
        pf = self._base(p)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.space_before = Pt(6)
        _setfont(p.add_run(text), 12, True)
        return p

    def center(self, text, bold=False):
        p = self.doc.add_paragraph()
        self._base(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _setfont(p.add_run(text), 12, bold)
        return p

    def body(self, text, bold=False):
        p = self.doc.add_paragraph()
        pf = self._base(p)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = Cm(1.25)
        _setfont(p.add_run(text), 12, bold)
        return p

    def quote(self, text):
        p = self.doc.add_paragraph()
        pf = self._base(p)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.left_indent = Cm(4)
        pf.first_line_indent = Cm(0)
        _setfont(p.add_run(text), 10, False)
        return p

    def action(self, text):
        p = self.doc.add_paragraph()
        pf = self._base(p)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = Cm(1.25)
        _setfont(p.add_run(text), 12, False)
        return p

    def blank(self):
        p = self.doc.add_paragraph()
        self._base(p)
        p.paragraph_format.first_line_indent = Cm(0)
        _setfont(p.add_run(""), 12)
        return p

    def image(self, img_path, width_cm=8.0, caption=None):
        """Insere uma imagem centralizada no corpo (prova/print)."""
        p = self.doc.add_paragraph()
        self._base(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run()
        run.add_picture(img_path, width=Cm(width_cm))
        if caption:
            c = self.doc.add_paragraph()
            self._base(c)
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.paragraph_format.first_line_indent = Cm(0)
            _setfont(c.add_run(caption), 10, False)
        return p

    def save(self, out):
        first = self.doc.paragraphs[0]
        if first.text.strip() == "":
            first._element.getparent().remove(first._element)
        self.doc.save(out)
        return out
